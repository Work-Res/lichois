from docx import Document
from docx.shared import Inches
from django.conf import settings
import os
from datetime import datetime


class WordDocumentService:
    """Responsible for generating a word document for assessment case summary.
    """
    def __init__(self, title, paragraphs, table_data, image_path=None):
        self.document = Document()
        self.title = title
        self.paragraphs = paragraphs
        self.table_data = table_data
        self.image_path = image_path

    def create_document(self):
        self._add_title()
        self._add_paragraphs()
        self._add_table()
        if self.image_path:
            self._add_image()
        return self.document

    def _add_title(self):
        self.document.add_heading(self.title, 0)

    def _add_paragraphs(self):
        for paragraph in self.paragraphs:
            self.document.add_paragraph(paragraph)

    def _add_table(self):
        if not self.table_data:
            return
        table = self.document.add_table(rows=len(self.table_data), cols=len(self.table_data[0]))
        table.style = 'Table Grid'
        for row_idx, row in enumerate(self.table_data):
            for col_idx, cell_text in enumerate(row):
                table.cell(row_idx, col_idx).text = cell_text

    def _add_image(self):
        if self.image_path:
            image_full_path = os.path.join(settings.MEDIA_ROOT, self.image_path)
            self.document.add_picture(image_full_path, width=Inches(1.25))

    def add_header(self, header_text, logo_path=None):
        for section in self.document.sections:
            header = section.header
            if logo_path:
                logo_full_path = os.path.join(settings.MEDIA_ROOT, logo_path)
                paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
                run = paragraph.add_run()
                run.add_picture(logo_full_path, width=Inches(1.25))
                paragraph.alignment = 1  # Center alignment
            paragraph = header.add_paragraph(header_text)
            paragraph.alignment = 1  # Center alignment

    def add_footer(self):
        for section in self.document.sections:
            footer = section.footer
            footer_paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()

            # Add date to the left
            date_run = footer_paragraph.add_run(datetime.now().strftime("%d/%m/%Y"))
            date_run.italic = True
            footer_paragraph.add_run(" " * 50)  # Adjust the number of spaces as needed

            # Add central text
            center_run = footer_paragraph.add_run("Assessment Case summary")
            center_run.italic = True

            # Add right text
            # footer_paragraph.add_run(" " * 50)  # Adjust the number of spaces as needed
            # right_run = footer_paragraph.add_run("LICHOIS")
            # right_run.italic = True

            footer_paragraph.alignment = 1  # Center alignment

    def save_document(self, file_path):
        self.document.save(file_path)
