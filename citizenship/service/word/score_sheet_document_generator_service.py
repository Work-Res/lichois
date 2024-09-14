import logging
from docx import Document
from docx.oxml.ns import nsdecls
from docx.oxml import OxmlElement

from docx2pdf import convert
from citizenship.exception.interview_score_sheet_error import (
    WordDocumentCreationError,
    PDFConversionError,
)

logger = logging.getLogger(__name__)


class ScoresheetDocumentGeneratorService:
    def __init__(self, data, word_path, pdf_path, template_path, context):
        """
        Initialize the service with data, template, and paths.

        Args:
            data (list): The data to populate the table.
            word_path (str): The path to save the populated Word document.
            pdf_path (str): The path to save the converted PDF document.
            template_path (str): The path to the saved Word document template.
        """
        self.data = data
        self.word_path = word_path
        self.pdf_path = pdf_path
        self.template_path = template_path
        self.context = context

    def create_word_document_from_template(self):
        """
        Create a Word document from a saved template and populate it with the provided data.
        """
        try:
            # Load the Word document template
            doc = Document(self.template_path)
            # Replace text in the document
            for paragraph in doc.paragraphs:
                for placeholder, replacement in self.context.items():
                    if placeholder in paragraph.text:
                        paragraph.text = paragraph.text.replace(placeholder, replacement)

            # Find the table in the template (assumes the first table is where you want to insert data)
            table = doc.add_table(rows=len(self.data), cols=4)  # Adjust if your template has multiple tables

            # Ensure that there are enough rows for the data, and add more rows if necessary
            # num_rows_needed = len(self.data)
            # num_existing_rows = len(table.rows)
            # for _ in range(num_rows_needed - num_existing_rows):
            #     table.add_row()

            # Populate the table with the provided data
            for row_idx, row_data in enumerate(self.data):
                row = table.rows[row_idx]
                for col_idx, cell_data in enumerate(row_data):
                    row.cells[col_idx].text = str(cell_data)
            # # Make table borders visible
            # for row in table.rows:
            #     for cell in row.cells:
            #         self.set_cell_border(cell)

            # Save the populated Word document to the specified path
            doc.save(self.word_path)
            logger.info(f"Word document created at {self.word_path}")
        except Exception as e:
            logger.exception(f"Error creating Word document from template: {e}")
            raise WordDocumentCreationError(f"Error creating Word document from template: {e}")

    def set_cell_border(self, cell):
        """
        Sets the borders for a given table cell.

        Args:
            cell: The table cell where borders need to be applied.
        """
        # Create the necessary elements for border formatting
        tc = cell._element
        tcPr = tc.get_or_add_tcPr()

        # Top border
        tcBorders = tcPr.get_or_add_tcBorders()
        top = OxmlElement('w:top')
        top.set(nsdecls('w'), 'single')
        top.set('w:sz', '4')
        top.set('w:color', '000000')
        tcBorders.append(top)

        # Left border
        left = OxmlElement('w:left')
        left.set(nsdecls('w'), 'single')
        left.set('w:sz', '4')
        left.set('w:color', '000000')
        tcBorders.append(left)

        # Bottom border
        bottom = OxmlElement('w:bottom')
        bottom.set(nsdecls('w'), 'single')
        bottom.set('w:sz', '4')
        bottom.set('w:color', '000000')
        tcBorders.append(bottom)

        # Right border
        right = OxmlElement('w:right')
        right.set(nsdecls('w'), 'single')
        right.set('w:sz', '4')
        right.set('w:color', '000000')
        tcBorders.append(right)
    def convert_to_pdf(self):
        """
        Convert the populated Word document to PDF.
        """
        try:
            # Convert the Word document to PDF
            convert(self.word_path, self.pdf_path)
            logger.info(f"PDF document created at {self.pdf_path}")
        except Exception as e:
            logger.exception(f"Error converting Word document to PDF: {e}")
            raise PDFConversionError(f"Error converting Word document to PDF: {e}")

    def create_and_convert(self):
        """
        Create the Word document from a template and convert it to PDF.
        """
        self.create_word_document_from_template()
        self.convert_to_pdf()
