import os
import shutil
from django.conf import settings
from django.test import TestCase, override_settings
from docx import Document

from app_information_requests.service.word import Generator
from citizenship.service.word import WordDocumentService


@override_settings(MEDIA_ROOT=os.path.join(settings.BASE_DIR, 'test_media'))
class WordDocumentServiceTest(TestCase):

    @classmethod
    def setUpClass(cls):
        super().setUpClass()
        cls.default_media_root = os.path.join(settings.BASE_DIR, 'media')
        cls.media_root = settings.MEDIA_ROOT
        os.makedirs(cls.media_root, exist_ok=True)

        # Copy all images from default MEDIA_ROOT to test_media directory
        if os.path.exists(cls.default_media_root):
            for root, _, files in os.walk(cls.default_media_root):
                for file in files:
                    if file.lower().endswith(('.png', '.jpg', '.jpeg', '.gif', '.bmp', '.tiff')):
                        src_file = os.path.join(root, file)
                        dest_file = os.path.join(cls.media_root, file)
                        shutil.copy(src_file, dest_file)
        else:
            # Create a dummy image file if no images exist in the default MEDIA_ROOT
            cls.sample_image_path = os.path.join(cls.media_root, 'logo.png')
            with open(cls.sample_image_path, 'wb') as f:
                f.write(b'\x00' * 1024)

    # @classmethod
    # def tearDownClass(cls):
    #     super().tearDownClass()
    #     if os.path.exists(cls.media_root):
    #         for root, dirs, files in os.walk(cls.media_root, topdown=False):
    #             for name in files:
    #                 os.remove(os.path.join(root, name))
    #             for name in dirs:
    #                 os.rmdir(os.path.join(root, name))
    #         os.rmdir(cls.media_root)

    def test_create_and_save_document_with_header_and_footer(self):
        title = 'Test Document'
        paragraphs = ['This is the first paragraph.', 'This is the second paragraph.']
        table_data = [
            ['Cell 1, 1', 'Cell 1, 2'],
            ['Cell 2, 1', 'Cell 2, 2']
        ]

        service = WordDocumentService(title, paragraphs, table_data, image_path='logo.png')
        document = service.create_document()
        service.add_header('', logo_path='logo.png')
        service.add_footer()

        # Save the document for viewing
        file_path = os.path.join(self.media_root, 'test_document.docx')
        service.save_document(file_path)

        # Check if the document is saved correctly
        self.assertTrue(os.path.exists(file_path))

        # Optionally, you can load the document and perform additional checks
        saved_document = Document(file_path)
        self.assertEqual(saved_document.paragraphs[0].text, title)
        self.assertEqual(saved_document.paragraphs[1].text, paragraphs[0])
        self.assertEqual(saved_document.paragraphs[2].text, paragraphs[1])

        tables = saved_document.tables
        self.assertEqual(len(tables), 1)
        table = tables[0]
        self.assertEqual(table.cell(0, 0).text, table_data[0][0])
        self.assertEqual(table.cell(0, 1).text, table_data[0][1])
        self.assertEqual(table.cell(1, 0).text, table_data[1][0])
        self.assertEqual(table.cell(1, 1).text, table_data[1][1])

        # Check if the header is correctly added
        for section in saved_document.sections:
            header = section.header
            self.assertTrue(any(p.text == '' for p in header.paragraphs))

    def test_add_title(self):
        title = 'Test Title'
        service = WordDocumentService(title, [], [])
        service._add_title()

        self.assertEqual(service.document.paragraphs[0].text, title)

    def test_add_paragraphs(self):
        paragraphs = ['First paragraph', 'Second paragraph']
        service = WordDocumentService('', paragraphs, [])
        service._add_paragraphs()

        self.assertEqual(service.document.paragraphs[0].text, paragraphs[0])
        self.assertEqual(service.document.paragraphs[1].text, paragraphs[1])

    def test_add_table(self):
        table_data = [
            ['Cell 1, 1', 'Cell 1, 2'],
            ['Cell 2, 1', 'Cell 2, 2']
        ]
        service = WordDocumentService('', [], table_data)
        service._add_table()

        tables = service.document.tables
        self.assertEqual(len(tables), 1)
        table = tables[0]
        self.assertEqual(table.cell(0, 0).text, table_data[0][0])
        self.assertEqual(table.cell(0, 1).text, table_data[0][1])
        self.assertEqual(table.cell(1, 0).text, table_data[1][0])
        self.assertEqual(table.cell(1, 1).text, table_data[1][1])

    def test_add_image(self):
        service = WordDocumentService('', [], [], image_path='logo.png')
        document = service.create_document()
        # Ensure the image is added (we'll check the presence of relationships as a proxy)
        rels = document.part.rels
        self.assertTrue(any(rel.target_ref.endswith('logo.png') for rel in rels.values()))

    def test_generator(self):
        logo_full_path = os.path.join(settings.MEDIA_ROOT, 'logo.png')
        output_path = os.path.join(self.media_root, 'test_document.docx')
        g = Generator(logo_path=logo_full_path)
        placeholders = {
            'full_name': 'John Doe',
            'checklist_request': 'Passport, Birth Certificate',
            'missing_information_request': 'Proof of Residence',
            'due_date': 'August 1, 2024',
            'contact_information': 'info@gov.bw',
            'officer_fullname': 'Jane Smith',
            'officer_position': 'Application Officer',
            'officer_contact_information': 'jane.smith@gov.bw'
        }
        g.create_request_letter(placeholders=placeholders, output_path=output_path)

