from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


class Generator:
    def __init__(self, logo_path):
        self.logo_path = logo_path

    def add_header(self, doc):
        # Access the header and add the logo
        section = doc.sections[0]
        header = section.header
        header_paragraph = header.paragraphs[0]
        header_paragraph.text = ""
        header_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Add logo to the header
        run = header_paragraph.add_run()
        run.add_picture(self.logo_path, width=Pt(60))  # Adjust the width as needed

    def add_footer(self, doc):
        # Access the footer and add some text
        section = doc.sections[0]
        footer = section.footer
        footer_paragraph = footer.paragraphs[0]
        footer_paragraph.text = "Confidential"
        footer_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    def create_request_letter(self, placeholders, output_path):

        doc = Document()
        self.add_header(doc)
        self.add_footer(doc)

        # Add the letter content with dynamic replacements
        doc.add_paragraph(f"\nDear {placeholders['full_name']},")

        doc.add_paragraph(
            'We hope this message finds you well. We are writing to inform you that we have reviewed '
            'your application and found that some important information is missing. To proceed with '
            'processing your application, we kindly request the following information:'
        )

        doc.add_paragraph(placeholders['checklist_request'])
        doc.add_paragraph(placeholders['missing_information_request'])

        p = doc.add_paragraph(
            f"Please provide this information by {placeholders['due_date']}. You can submit the information via "
            "[preferred method, e.g., reply to this email, upload to our gov online portal]."
        )
        p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

        doc.add_paragraph(
            f"If you have any questions or need assistance, please do not hesitate to contact us at "
            f"{placeholders['contact_information']}."
        )

        doc.add_paragraph('Thank you for your prompt attention to this matter.\n')

        doc.add_paragraph('Best regards,\n')

        doc.add_paragraph(f"{placeholders['officer_fullname']}\n{placeholders['officer_position']}\n{placeholders['officer_contact_information']}")

        # Save the document
        doc.save(output_path)
