import os
from django.conf import settings
from django.test import TestCase
from django.test import TestCase, override_settings

from docx import Document

from app_information_requests.service.word import WordDocumentTemplateService, Generator


@override_settings(MEDIA_ROOT=os.path.join(settings.BASE_DIR, 'test_media'))
class WordDocumentTemplateServiceTest(TestCase):

    @classmethod
    def setUpClass(cls):
        super().setUpClass()
        cls.default_media_root = os.path.join(settings.BASE_DIR, 'media')
        cls.media_root = settings.MEDIA_ROOT
        os.makedirs(cls.media_root, exist_ok=True)

    @classmethod
    def setUpClass(cls):
        super().setUpClass()
        cls.default_media_root = os.path.join(settings.BASE_DIR, 'media')
        cls.media_root = settings.MEDIA_ROOT
        os.makedirs(cls.media_root, exist_ok=True)
        # Create a sample template.docx file with placeholders for testing
        cls.template_path = 'template.docx'
        cls.output_path = 'output.docx'
        doc = Document()
        doc.add_paragraph('Dear {{ name }},')
        doc.add_paragraph('Your application for {{ position }} has been approved.')
        doc.add_paragraph('Best regards,')
        doc.add_paragraph('{{ company_name }}')
        doc.save(cls.template_path)

    # @classmethod
    # def tearDownClass(cls):
    #     super().tearDownClass()
    #     # Clean up the test files
    #     os.remove(cls.template_path)
    #     if os.path.exists(cls.output_path):
    #         os.remove(cls.output_path)

    def test_replace_placeholders(self):
        service = WordDocumentTemplateService(self.template_path)
        context = {
            'name': 'John Doe',
            'position': 'Software Engineer',
            'company_name': 'TechCorp'
        }
        document = service.replace_placeholders(context)
        service.save_document(document, self.output_path)

        # Load the output document to check if placeholders are replaced
        output_doc = Document(self.output_path)
        paragraphs = [p.text for p in output_doc.paragraphs]

        self.assertIn('Dear John Doe,', paragraphs)
        self.assertIn('Your application for Software Engineer has been approved.', paragraphs)
        self.assertIn('Best regards,', paragraphs)
        self.assertIn('TechCorp', paragraphs)

    def test_generator(self):
        service = Generator('path/to/logo.png')
        placeholders = {
            'full_name': 'John Doe',
            'checklist_request': 'Passport, Birth Certificate',
            'missing_information_request': 'Proof of Residence',
            'due_date': 'August 1, 2024',
            'contact_information': 'info@gov.bw',
            'officer_fullname': 'Jane Smith',
            'officer_position': 'Application Officer',
            'officer_contact_information': 'jane.smith@gov.bw'
        }
        service.create_request_letter(placeholders, 'request_letter.docx')


    def test_replace_placeholders_new_template(self):
        self.template_path = "sample_information_request.docx"
        self.output_path = "sample_information_request_1.docx"
        service = WordDocumentTemplateService(self.template_path)
        context = {
            'full_name': 'Tshepiso Setsiba',
            'officer_fullname': 'Moffat Motlhanka',
            'position': 'Senior Software Engineer',
            'officer_contact_information': '+267 73424507/tsetsiba@sample.com',
            'missing_information_request': 'A request for proposal (RFP) is a project announcement posted publicly by an organization indicating that bids for contractors to complete the project are sought.'
        }
        document = service.replace_placeholders(context)
        service.save_document(document, self.output_path)

        # Load the output document to check if placeholders are replaced
        output_doc = Document(self.output_path)
        paragraphs = [p.text for p in output_doc.paragraphs]

        # self.assertIn('Dear John Doe,', paragraphs)
        # self.assertIn('Your application for Software Engineer has been approved.', paragraphs)
        # self.assertIn('Best regards,', paragraphs)
        # self.assertIn('TechCorp', paragraphs)
