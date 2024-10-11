import docx
from docx2pdf import convert
import logging
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


class GenerateGazettePDFService:
    def __init__(self, data, word_file_path, pdf_file_path):
        if not isinstance(data, list) or not all(isinstance(row, list) for row in data):
            raise ValueError("Data must be a list of lists.")
        self.data = data
        self.word_file_path = word_file_path
        self.pdf_file_path = pdf_file_path


    def create_word_document(self):
        logger = logging.getLogger(__name__)
        try:

            doc = docx.Document()

            citizenship_act_heading = doc.add_heading('CITIZENSHIP ACT', level=1)
            citizenship_act_heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            cap_paragraph = doc.add_paragraph('(Cap. 01:01)')
            cap_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            sub_heading_1 = doc.add_heading('Application for Certificate of Naturalization', level=2)
            sub_heading_1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            notice_text_1 = (
                "NOTICE IS HEREBY GIVEN that in accordance with section 13 (3) and (4) of the Citizenship Act, "
                "each of the persons whose names and addresses are listed below have applied to the Minister of Labour "
                "and Home Affairs for a Certificate of Naturalization. Any person who objects to any of the applications "
                "should notify the Minister, in writing, at Private Bag 002, Gaborone of his or her objection and the "
                "grounds thereof within thirty (30) days of publication of this Notice."
            )
            doc.add_paragraph(notice_text_1)

            sub_heading_2 = doc.add_heading('Kopo ya Boagedi', level=2)
            sub_heading_2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            notice_text_2 = (
                "BATHO BA MAINA A, ba ikopela boagedi jwa Botswana mo go Tona wa Lephata la Pereko le Selegae. "
                "Fa go na le mongwe yo o ka bong a le kgatlhanong le kopo ya mongwe wa bakopi ba, o ka kwalela kwa go "
                "Tona wa tsa Pereko le Selegae, a tlhalosa mabaka a ka one a leng kgatlhanong le kopo eo. "
                "Dilwalo tsa go nna jalo di goroge kwa go Tona ko Private Bag 002, Gaborone mo malatsing a le masome "
                "a mararo (30) go simolola tsatsi le kitsiso e duleng ka lone."
            )
            doc.add_paragraph(notice_text_2)

            # Create a table
            table = doc.add_table(rows=1, cols=len(self.data[0]))
            table.style = 'Table Grid'

            # Add header row
            hdr_cells = table.rows[0].cells
            for i, heading in enumerate(self.data[0]):
                hdr_cells[i].text = heading

            # Add data rows
            for row in self.data[1:]:
                row_cells = table.add_row().cells
                for i, cell in enumerate(row):
                    row_cells[i].text = str(cell)

            # Save the document
            doc.save(self.word_file_path)
        except Exception as e:
            logger.error(f"Error creating Word document: {e}")

    def convert_to_pdf(self):
        # Convert the Word document to PDF
        convert(self.word_file_path, self.pdf_file_path)

    def generate_pdf(self):
        self.create_word_document()
        self.convert_to_pdf()
