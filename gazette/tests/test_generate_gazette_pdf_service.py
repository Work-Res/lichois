import os
import docx

from django.conf import settings
from spire.doc import *
from spire.doc.common import *

from datetime import datetime


from gazette.service import GenerateGazettePDFService, PrepareGazetteForDownload

from .base_setup import BaseSetup
from ..models import Batch, BatchApplication


class TestGenerateGazettePDFService(BaseSetup):

    def setUp(self):
        super().setUp()
        self.data = [
            ["ID", "Fullname", "Location", "Document Number"],
            ["1", "Test test", "Ranaka", "RZn0001"],
            ["2", "Bryce Sets", "Kanye", "R2000001"],
            ["3", "Laim Kgathi", "Moshupa", "R4000333"]
        ]

        self.create_apps_villages()

        batch = Batch.objects.get(
            status="IN_PROGRESS"
        )
        batch_apps = BatchApplication.objects.filter(
            batch=batch
        )
        now = datetime.now()
        formatted_date = now.strftime("%Y-%m-%d")

        prepare_download_data = PrepareGazetteForDownload(batch_applications=batch_apps)
        data = prepare_download_data.prepared_data()
        self.word_file_path = os.path.join(settings.MEDIA_ROOT, f"gazette_{formatted_date}.docx")
        self.pdf_file_path = os.path.join(settings.MEDIA_ROOT, f"test_applications_{formatted_date}.pdf")
        self.service = GenerateGazettePDFService(data, self.word_file_path, self.pdf_file_path)
    # def tearDown(self):
    #     # Clean up temporary files
    #     if os.path.exists(self.word_file_path):
    #         os.remove(self.word_file_path)
    #     if os.path.exists(self.pdf_file_path):
    #         os.remove(self.pdf_file_path)

    def test_create_word_document(self):
        self.service.create_word_document()
        # Ensure the Word document is created
        self.assertTrue(os.path.exists(self.word_file_path))

        # # Additional checks for document content
        # doc = docx.Document(self.word_file_path)
        # self.assertEqual(doc.paragraphs[0].text, "Gazette List")
        # table = doc.tables[0]
        # self.assertEqual(len(table.rows), 4)  # 1 header row + 3 data rows
        # self.assertEqual(table.cell(0, 0).text, "ID")
        # self.assertEqual(table.cell(1, 0).text, "1")
        # self.assertEqual(table.cell(2, 1).text, "Bryce Sets")
        # self.assertEqual(table.cell(3, 2).text, "Moshupa")
        # # Create a Document object
        # document = Document()
        # # Load a Word DOCX file
        # document.LoadFromFile(self.word_file_path)
        # # Or load a Word DOC file
        # # document.LoadFromFile("Sample.doc")
        #
        # # Save the file to a PDF file
        # document.SaveToFile("WordToPdf.pdf", FileFormat.PDF)
        # document.Close()

    def test_convert_to_pdf(self):
        self.service.create_word_document()
        self.service.convert_to_pdf()
        # Ensure the PDF document is created
        self.assertTrue(os.path.exists(self.pdf_file_path))

    def test_generate_pdf(self):
        self.service.generate_pdf()
        # Ensure both Word and PDF documents are created
        self.assertTrue(os.path.exists(self.word_file_path))
        self.assertTrue(os.path.exists(self.pdf_file_path))
