from docx import Document
from docx2pdf import convert


class WordDocumentTemplateService:

    def __init__(self, template_path):
        self.template_path = template_path

    def replace_placeholders(self, context):
        # Load the template document
        doc = Document(self.template_path)

        # Replace placeholders in paragraphs
        for paragraph in doc.paragraphs:
            for key, value in context.items():
                if f'{{{{ {key} }}}}' in paragraph.text:
                    paragraph.text = paragraph.text.replace(f'{{{{ {key} }}}}', value)

        # Replace placeholders in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for key, value in context.items():
                        if f'{{{{ {key} }}}}' in cell.text:
                            cell.text = cell.text.replace(f'{{{{ {key} }}}}', value)
        # Replace placeholders in headers and footers
        self._replace_in_headers_footers(doc, context)

        return doc

    def _replace_in_headers_footers(self, doc, context):
        # Iterate over all sections
        for section in doc.sections:
            # Replace placeholders in headers
            for header in section.header.paragraphs:
                for key, value in context.items():
                    if f'{{{{ {key} }}}}' in header.text:
                        header.text = header.text.replace(f'{{{{ {key} }}}}', value)

            # Replace placeholders in footers
            for footer in section.footer.paragraphs:
                for key, value in context.items():
                    if f'{{{{ {key} }}}}' in footer.text:
                        footer.text = footer.text.replace(f'{{{{ {key} }}}}', value)

    def save_document(self, doc, output_path):
        doc.save(output_path)

    def convert_to_pdf(self, doc_path, pdf_output_path):
        # Convert the Word document to a PDF using docx2pdf
        if pdf_output_path:
            convert(doc_path, pdf_output_path)
            print(f"PDF saved to: {pdf_output_path}")
